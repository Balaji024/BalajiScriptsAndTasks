import docx

d = docx.Document('H:\Resume_Format.docx')
print(d.paragraphs)
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)
p = d.paragraphs[1]
print(p.runs[0].text)
print(p.runs[1].text)
print(p.style)
#p.style = 'Title'
#d.save('H:\Resume_Format.docx')
d1 = docx.Document()
d1.add_paragraph("Synopsis")
d1.add_paragraph("Introduction")
d1.save('H:\mysample.docx')

def getText(filename):

  doc = docx.Document(filename)
  fulltext = []
  for para in doc.paragraphs:
      fulltext.append(para.text)
  return '\n'.join(fulltext)

print(getText('H:\mysample.docx'))





